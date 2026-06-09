# -*- coding: utf-8 -*-
"""Insere a seção Resultados e Discussão no TCC II (1).docx"""

import copy
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

DOC_PATH = "TCC II (1).docx"
OUT_PATH = "TCC II (1).docx"

doc = Document(DOC_PATH)
body = doc.element.body


# ── helpers ──────────────────────────────────────────────────────────────────

def get_body_children():
    return list(body)


def remove_children_range(start_idx, end_idx_exclusive):
    """Remove body children entre start e end (exclusive)."""
    children = list(body)
    to_remove = children[start_idx:end_idx_exclusive]
    for child in to_remove:
        body.remove(child)


def make_paragraph_xml(text, bold=False, italic=False, align=None):
    """Cria elemento <w:p> com estilo 'normal' e texto fornecido."""
    p = OxmlElement("w:p")

    pPr = OxmlElement("w:pPr")
    pStyle = OxmlElement("w:pStyle")
    pStyle.set(qn("w:val"), "normal")
    pPr.append(pStyle)
    if align:
        jc = OxmlElement("w:jc")
        jc.set(qn("w:val"), align)
        pPr.append(jc)
    p.append(pPr)

    r = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    if bold:
        b = OxmlElement("w:b")
        rPr.append(b)
    if italic:
        i = OxmlElement("w:i")
        rPr.append(i)
    r.append(rPr)

    t = OxmlElement("w:t")
    t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    t.text = text
    r.append(t)
    p.append(r)
    return p


def insert_elements_after(anchor_el, elements):
    """Insere lista de elementos após anchor_el, em ordem."""
    # Inserir em ordem reversa todos após o mesmo anchor
    for el in reversed(elements):
        anchor_el.addnext(el)


def make_table_xml(headers, rows, caption=None, fonte=None):
    """Retorna lista de elementos XML: [caption_p, tbl, fonte_p]"""
    result = []

    if caption:
        result.append(make_paragraph_xml(caption, bold=True))

    # Cria a tabela via python-docx no documento temporário
    tmp_doc = Document()
    tbl = tmp_doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = "Table Grid"

    # Cabeçalho
    hdr_cells = tbl.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for run in hdr_cells[i].paragraphs[0].runs:
            run.bold = True

    # Linhas de dados
    for r_idx, row_data in enumerate(rows):
        row_cells = tbl.rows[r_idx + 1].cells
        for c_idx, cell_text in enumerate(row_data):
            row_cells[c_idx].text = cell_text

    tbl_el = copy.deepcopy(tbl._tbl)
    result.append(tbl_el)

    if fonte:
        result.append(make_paragraph_xml(fonte, italic=True))

    return result


# ── Localizar os índices dos elementos no body ────────────────────────────────

children = get_body_children()

idx_resultados = None
idx_consideracoes = None

for i, child in enumerate(children):
    tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
    if tag == "p":
        text = "".join(
            r.text or ""
            for r in child.iter(qn("w:t"))
        )
        if "RESULTADOS E DISCUSS" in text.upper():
            idx_resultados = i
        if "CONSIDERA" in text.upper() and "FINAIS" in text.upper():
            idx_consideracoes = i
            break

print(f"idx_resultados={idx_resultados}, idx_consideracoes={idx_consideracoes}")

# Elementos placeholder a remover: tudo entre os dois marcadores
remove_children_range(idx_resultados + 1, idx_consideracoes)

# Recarrega lista após remoção
children = get_body_children()
anchor = children[idx_resultados]  # "4. RESULTADOS E DISCUSSÃO"


# ── Conteúdo da seção ─────────────────────────────────────────────────────────

elementos = []

# ── Introdução ──
elementos.append(make_paragraph_xml(
    "A presente seção expõe os resultados obtidos nas etapas de treinamento e avaliação "
    "do pipeline híbrido proposto, procedendo à análise das métricas extraídas do subconjunto "
    "de teste e ao confronto dos achados com o estado da arte revisado na Seção 2."
))

# ── 4.1 ──
elementos.append(make_paragraph_xml("4.1 Métricas de Avaliação do Modelo Multiclasse (ResNet50)", bold=True))

elementos.append(make_paragraph_xml(
    "O modelo de classificação multiclasse foi avaliado sobre 6.563 imagens do subconjunto "
    "de teste (hold-out de 20%), resultantes do particionamento estratificado descrito na "
    "Seção 3.1. O Quadro 2 consolida as métricas de desempenho obtidas por classe de criticidade."
))

headers_q2 = ["Classe", "Precisão", "Recall", "F1-Score", "Suporte"]
rows_q2 = [
    ["Crítica",      "5,41%",  "100,00%", "10,26%", "4"],
    ["Semi-Crítica", "60,66%", "18,29%",  "28,10%", "3.702"],
    ["Não Crítica",  "45,00%", "84,63%",  "58,76%", "2.857"],
    ["Acurácia Global", "", "", "47,22%", ""],
    ["F1-Macro",        "", "", "32,37%", ""],
    ["F1-Weighted",     "", "", "41,44%", ""],
]
elementos += make_table_xml(
    headers_q2, rows_q2,
    caption="Quadro 2 — Métricas de desempenho do classificador ResNet50 multiclasse por classe de criticidade.",
    fonte="Fonte: Elaborado pelo autor (2026)."
)

elementos.append(make_paragraph_xml(
    "A acurácia global de 47,22% e o F1-Macro de 32,37% refletem, primariamente, o severo "
    "desbalanceamento de classes inerente ao dataset rotulado: a classe Crítica conta com apenas "
    "4 amostras no subconjunto de teste, insuficientes para qualquer inferência estatística robusta. "
    "O comportamento desta classe merece análise particularizada: o Recall de 100% indica que o "
    "modelo identificou corretamente todas as instâncias críticas presentes, ao custo de uma "
    "Precisão de 5,41%, reflexo da elevada taxa de falsos positivos gerada para essa categoria. "
    "Tal resultado é coerente com a lógica da Focal Loss (LIN et al., 2017) adotada durante o "
    "treinamento — o fator modulador concentra o gradiente nos exemplos de difícil aprendizado, "
    "o que, para uma classe com apenas 24 amostras de treino, equivale a induzir o modelo a nunca "
    "omitir uma instância crítica, ainda que à custa de falsos alarmes. Do ponto de vista da "
    "segurança estrutural, esse comportamento é o mais conservador e, portanto, o mais desejável: "
    "em inspeções de estruturas de concreto, a omissão de uma fissura crítica representa risco "
    "irreversível, ao passo que um falso positivo pode ser verificado in loco pelo engenheiro responsável."
))

elementos.append(make_paragraph_xml(
    "O desempenho da classe Semi-Crítica apresenta Precisão de 60,66% com Recall de apenas "
    "18,29%, indicando que a maioria das amostras desta categoria foi classificada erroneamente "
    "como Não Crítica. Tal tendência reflete a proximidade morfológica entre as duas classes — "
    "os limiares de 0,2 mm e 1,0 mm estabelecidos pela ABNT NBR 6118:2014 constituem fronteiras "
    "geométricas de difícil discriminação visual quando as imagens não apresentam diferenciação "
    "tonal nítida entre as categorias. A classe Não Crítica, por ser majoritária (66,90% do "
    "dataset), obteve o melhor desempenho relativo, com F1-Score de 58,76%, resultado esperado "
    "haja vista a disponibilidade de 37.462 amostras de treino para essa categoria."
))

# ── 4.2 ──
elementos.append(make_paragraph_xml("4.2 Métricas de Avaliação do Modelo Binário (Intervenção vs. Não Crítica)", bold=True))

elementos.append(make_paragraph_xml(
    "Complementarmente ao classificador multiclasse, procedeu-se ao treinamento e à avaliação de "
    "um modelo binário, destinado a distinguir imagens que demandam intervenção estrutural — "
    "classes Crítica e Semi-Crítica agrupadas sob o rótulo Intervenção — daquelas classificadas "
    "como Não Críticas. O modelo foi avaliado sobre 6.564 imagens do subconjunto de teste, "
    "correspondentes à partição de 80%/20% com semente aleatória 42. O Quadro 3 apresenta os resultados obtidos."
))

headers_q3 = ["Classe", "Precisão", "Recall", "F1-Score", "Suporte"]
rows_q3 = [
    ["Intervenção",     "58,95%", "53,82%", "56,27%", "3.707"],
    ["Não Crítica",     "46,16%", "51,38%", "48,63%", "2.857"],
    ["Acurácia Global", "",       "",       "52,76%", ""],
    ["F1-Macro",        "",       "",       "52,45%", ""],
    ["F1-Weighted",     "",       "",       "52,95%", ""],
]
elementos += make_table_xml(
    headers_q3, rows_q3,
    caption="Quadro 3 — Métricas de desempenho do classificador binário ResNet50 (Intervenção vs. Não Crítica).",
    fonte="Fonte: Elaborado pelo autor (2026)."
)

elementos.append(make_paragraph_xml(
    "A acurácia de 52,76% no cenário binário, embora modesta em termos absolutos, deve ser "
    "contextualizada à luz das condições de treinamento: o modelo foi submetido a um número "
    "restrito de épocas, em decorrência de limitações de infraestrutura computacional, conforme "
    "detalhado na Seção 4.5. A matriz de confusão registra 1.995 verdadeiros positivos para a "
    "classe Intervenção, contra 1.712 falsos negativos — indicando que aproximadamente 46% dos "
    "casos que demandariam atenção foram classificados incorretamente como Não Críticos. Em "
    "cenário de produção, esse resultado é mitigado pela sobrescrita de segurança estrutural: "
    "qualquer imagem para a qual o classificador físico detecte largura superior a 1,0 mm tem "
    "sua classificação forçada para Crítica, independentemente da saída da rede neural."
))

# ── 4.3 ──
elementos.append(make_paragraph_xml("4.3 Análise das Curvas de Aprendizado e Convergência", bold=True))

elementos.append(make_paragraph_xml(
    "O treinamento do classificador ResNet50 multiclasse foi conduzido com a estratégia de "
    "descongelamento progressivo (Staged Fine-Tuning), descrita na Seção 3.4.2. Nas épocas "
    "iniciais, o backbone permaneceu congelado e apenas a cabeça totalmente conectada foi "
    "atualizada, possibilitando que os pesos pré-treinados no ImageNet fossem preservados enquanto "
    "a nova camada de classificação convergisse para as distribuições específicas do SDNET2018. "
    "Nas épocas subsequentes, as camadas layer3 e layer4 foram descongeladas com taxa de "
    "aprendizado diferenciada (LR_backbone = LR_head x 0,1), estratégia que previne a degradação "
    "catastrófica das representações de baixo nível aprendidas no pré-treinamento "
    "(GOODFELLOW et al., 2016)."
))

elementos.append(make_paragraph_xml(
    "O checkpoint selecionado ao término do treinamento foi determinado pelo maior F1-Score Macro "
    "observado na partição de validação — critério adotado em substituição à minimização da perda "
    "de validação, tendo em vista que o F1-Macro pondera igualmente as três classes, forçando o "
    "modelo a apresentar desempenho relevante também na categoria minoritária Crítica."
))

# ── 4.4 ──
elementos.append(make_paragraph_xml("4.4 Comportamento do Pipeline Híbrido e Complementaridade dos Componentes", bold=True))

elementos.append(make_paragraph_xml(
    "Os resultados isolados do componente de Aprendizado Profundo devem ser interpretados em "
    "conjunto com o papel desempenhado pelo componente físico do pipeline híbrido. Quando a "
    "confiança do classificador neural situa-se abaixo dos limiares individuais por classe "
    "(Crítica = 0,65; Semi-Crítica = 0,55; Não Crítica = 0,50), a decisão é delegada ao "
    "classificador por limiares físicos, fundamentado na largura milimétrica estimada pela "
    "Transformada de Distância. Essa complementaridade é particularmente relevante para a classe "
    "Crítica: imagens com largura mediana superior a 1,0 mm receberam confiança média de 0,91 "
    "por parte do classificador físico, com valores mínimos observados de 0,76 — evidenciando "
    "que o componente de medição geométrica opera com elevada segurança precisamente nas fissuras "
    "mais severas, que são aquelas em que o componente neural apresenta maior instabilidade "
    "estatística em decorrência da subrepresentação da classe no dataset de treinamento."
))

# ── 4.5 ──
elementos.append(make_paragraph_xml("4.5 Limitações do Treinamento", bold=True))

elementos.append(make_paragraph_xml(
    "A análise comparativa entre as métricas de treino e teste evidencia que o modelo operou sob "
    "condições limitantes que comprometeram parcialmente sua capacidade de generalização. O principal "
    "fator restritivo foi o número reduzido de épocas de treinamento, imposto pelas limitações de "
    "infraestrutura computacional disponível — o processamento foi realizado integralmente em CPU "
    "(Unidade Central de Processamento), sem aceleração por GPU (Unidade de Processamento Gráfico), "
    "o que eleva consideravelmente o custo temporal por época. O desbalanceamento severo de classes "
    "constitui o segundo vetor de degradação de desempenho: a proporção de 0,04% para a classe "
    "Crítica — correspondente a 24 amostras de treino em um dataset de aproximadamente 55.000 "
    "imagens — representa um caso extremo de escassez de dados que nenhuma estratégia de otimização, "
    "seja Focal Loss, limiares por classe ou descongelamento progressivo, é capaz de compensar "
    "integralmente sem a incorporação de novas amostras reais."
))

# ── 4.6 ──
elementos.append(make_paragraph_xml("4.6 Comparação com Trabalhos Relacionados", bold=True))

elementos.append(make_paragraph_xml(
    "Os resultados obtidos pelo pipeline híbrido inserem-se num espectro de desempenho coerente "
    "com o estado da arte revisado na Seção 2, consideradas as especificidades da tarefa de "
    "classificação de criticidade — problema substancialmente mais complexo do que a detecção "
    "binária convencional abordada pela maior parte da literatura."
))

elementos.append(make_paragraph_xml(
    "Ali et al. (2021) reportaram acurácias entre 85% e 97% para a tarefa de detecção binária "
    "(presença ou ausência de fissura) em estruturas de concreto, empregando arquiteturas como "
    "VGG16 e ResNet sobre datasets com distribuições balanceadas. Contudo, o problema abordado "
    "por esses autores é qualitativamente distinto do proposto no presente estudo: a detecção "
    "binária opera sobre uma fronteira de decisão única e bem definida, ao passo que a "
    "classificação de criticidade em três níveis — com limiares geométricos da ordem de décimos "
    "de milímetro estabelecidos pela ABNT NBR 6118:2014 — impõe uma dificuldade discriminativa "
    "de outra ordem de grandeza."
))

elementos.append(make_paragraph_xml(
    "Philip et al. (2023), em estudo comparativo de Aprendizado por Transferência para a "
    "detecção de fissuras em paredes de concreto, obtiveram F1-Scores de até 94% com a "
    "arquitetura ResNet50 em datasets balanceados de natureza binária. König et al. (2022), "
    "em revisão abrangente de métodos de Aprendizado Profundo para segmentação e quantificação "
    "de fissuras, destacam que a estimativa de largura milimétrica representa uma das fronteiras "
    "mais abertas da área, com os melhores modelos reportando erros médios da ordem de 0,1 mm — "
    "margem que coincide com a faixa de incerteza crítica para os limiares normativos da "
    "ABNT NBR 6118:2014."
))

elementos.append(make_paragraph_xml(
    "O diferencial metodológico do presente trabalho reside não na acurácia isolada do componente "
    "neural, mas na integração híbrida entre a inferência semântica da ResNet50 e a validação "
    "física fundamentada na Transformada de Distância. Essa arquitetura permite que o sistema "
    "opere com confiança mensurável mesmo nos casos em que a rede neural apresenta baixa certeza "
    "— mecanismo ausente nos trabalhos comparados. Ademais, a incorporação da sobrescrita de "
    "segurança estrutural representa uma contribuição de natureza aplicada: ao assegurar que "
    "nenhuma fissura fisicamente Crítica seja classificada de forma menos severa, o sistema "
    "prioriza a integridade da inspeção sobre a otimização estatística das métricas, "
    "alinhando-se aos requisitos normativos da engenharia diagnóstica."
))

# ── 4.7 ──
elementos.append(make_paragraph_xml("4.7 Limitações Identificadas", bold=True))

elementos.append(make_paragraph_xml(
    "A primeira e mais determinante limitação é a escassez de amostras da classe Crítica no "
    "dataset SDNET2018. Com apenas 24 instâncias para treino e 4 para teste, nenhuma métrica "
    "calculada para essa classe pode ser considerada estatisticamente representativa. A construção "
    "de um modelo confiável para a detecção de fissuras macroscopicamente severas requer, "
    "necessariamente, a coleta e a anotação de novos dados reais — preferencialmente obtidos em "
    "inspeções de campo realizadas com acompanhamento de engenheiros estruturais e rastreabilidade normativa."
))

elementos.append(make_paragraph_xml(
    "A segunda limitação refere-se ao domínio dos dados de treinamento. O SDNET2018 é composto "
    "por sub-recortes padronizados de superfícies de concreto, capturados em condições controladas. "
    "A aplicação do modelo sobre imagens de campo reais — com variação de iluminação, perspectiva, "
    "texturas naturais e obstruções — configura um problema de transferência de domínio "
    "(domain shift) que tende a degradar as métricas reportadas, demandando estratégias de "
    "adaptação de domínio para uso em produção."
))

elementos.append(make_paragraph_xml(
    "A terceira limitação diz respeito ao processo de rotulagem automatizada. Os rótulos do "
    "dataset foram gerados pelo próprio algoritmo morfométrico, o que cria uma dependência "
    "circular: o classificador neural foi treinado para reproduzir as decisões do classificador "
    "físico, potencialmente herdando os vieses e os erros sistemáticos inerentes ao rotulador "
    "automático. A validação dos rótulos por especialistas em engenharia estrutural constitui, "
    "portanto, requisito para a consolidação do sistema em ambiente de produção."
))

elementos.append(make_paragraph_xml(
    "A quarta limitação concerne à explicabilidade das predições. Em contextos de engenharia "
    "diagnóstica, o engenheiro responsável necessita não apenas da classificação final, mas da "
    "fundamentação objetiva da decisão. O pipeline atual disponibiliza a confiança numérica da "
    "ResNet50 e o valor de largura milimétrica como elementos de rastreabilidade, não "
    "incorporando, contudo, técnicas formais de interpretabilidade como o Mapeamento de Ativação "
    "de Classe ponderado por Gradiente (Grad-CAM, do inglês Gradient-weighted Class Activation "
    "Mapping) ou SHAP (SHapley Additive exPlanations), cuja integração é identificada como "
    "direção prioritária para versões futuras do sistema."
))


# ── Inserção no documento ─────────────────────────────────────────────────────

insert_elements_after(anchor, elementos)

doc.save(OUT_PATH)
print(f"Documento salvo: {OUT_PATH}")
print(f"Total de elementos inseridos: {len(elementos)}")
